import streamlit as st
import os
import io
import pandas as pd
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
import nltk
from transformers import pipeline
from sklearn.feature_extraction.text import TfidfVectorizer
from keybert import KeyBERT
from nltk.tokenize import sent_tokenize
import re

# Download necessary NLTK datasets
nltk.download('punkt')  # sentence tokenization
nltk.download('stopwords')

# Initialize Hugging Face summarizer
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# Initialize KeyBERT
kw_model = KeyBERT()

class EnhancedSummarizer:
    def __init__(self, summarizer):
        self.summarizer = summarizer

    def summarize(self, text: str) -> str:
        max_length = 1024  # Adjust based on model limits
        if len(text) > max_length:
            truncated_text = text[:max_length // 2] + text[-max_length // 2:]
        else:
            truncated_text = text

        summary = self.summarizer(
            truncated_text,
            max_length=150,
            min_length=50,
            do_sample=False
        )[0]['summary_text']
        return summary

    def extract_keywords(self, text: str, num_keywords: int = 5) -> list:
        keywords = kw_model.extract_keywords(
            text,
            vectorizer=TfidfVectorizer(stop_words="english"),
            top_n=num_keywords
        )
        return [keyword[0] for keyword in keywords]

    def generate_highlights(self, text: str, num_highlights: int = 5) -> list:
        """
        Generate meaningful highlights by combining TF-IDF scores with keyword relevance.
        """
        # Tokenize the text into sentences
        sentences = sent_tokenize(text)

        # Extract keywords from the text
        keywords = self.extract_keywords(text, num_keywords=10)  # Extract more keywords for better coverage

        # Use TF-IDF to score sentences
        vectorizer = TfidfVectorizer(stop_words="english")
        sentence_vectors = vectorizer.fit_transform(sentences)
        sentence_scores = sentence_vectors.sum(axis=1).A1  # Convert matrix to 1D array

        # Adjust scores based on keyword relevance
        keyword_scores = [
            sum(1 for word in keywords if word.lower() in sentence.lower()) for sentence in sentences
        ]

        # Combine TF-IDF and keyword relevance scores
        combined_scores = [
            tfidf_score + (2 * keyword_score)  # Weight keyword relevance higher than TF-IDF
            for tfidf_score, keyword_score in zip(sentence_scores, keyword_scores)
        ]

        # Rank sentences based on combined scores
        ranked_sentences = [sentences[i] for i in sorted(range(len(combined_scores)), key=lambda x: combined_scores[x], reverse=True)]

        # Filter out irrelevant or poorly formatted sentences
        filtered_sentences = [
            sentence for sentence in ranked_sentences
            if len(sentence.split()) >= 8  # Filter out very short sentences
            and not re.search(r'http\S+|www\S+', sentence)  # Remove sentences with URLs
            and not re.search(r'[^a-zA-Z0-9\s.,;:!?]', sentence)  # Remove sentences with special characters
            and not re.search(r'\b[A-Z]{2,}\b', sentence)  # Remove sentences with all-caps words
            and not re.search(r'\d+\.\s+', sentence)  # Remove sentences with numbers (e.g., "2.3 Structure")
        ]

        # If no sentences pass the filter, use top-ranked sentences without filtering
        if not filtered_sentences:
            filtered_sentences = ranked_sentences[:num_highlights]

        # Select the top sentences as highlights
        highlights = [
            f"- 📊 {sentence.strip()}"
            for sentence in filtered_sentences[:num_highlights]
        ]

        return highlights


enhanced_summarizer = EnhancedSummarizer(summarizer)

def extract_text_from_docx(docx_path):
    document = Document(docx_path)
    text = ""
    for para in document.paragraphs:
        text += para.text + "\n"
    return text

def clean_text(text: str) -> str:
    text = text.replace("\n", " ").strip()
    text = " ".join(text.split())
    return text

def chunk_text(text, chunk_size=1000):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i+chunk_size])
        chunks.append(chunk)
    return chunks

def generate_structured_output(text):
    cleaned_text = clean_text(text)
    text_chunks = chunk_text(cleaned_text)

    summaries = [enhanced_summarizer.summarize(chunk) for chunk in text_chunks if chunk.strip()]
    full_summary = " ".join(summaries)

    summary_sentences = [s for s in full_summary.split(". ") if not re.search(r'\d+', s)]
    limited_summary = ". ".join(summary_sentences[:5]) + "."

    highlights = enhanced_summarizer.generate_highlights(cleaned_text, num_highlights=7)

    structured_output = (
        "# Summary\n\n" +
        f"{limited_summary}\n\n" +
        "### Highlights\n" +
        "\n".join(highlights) +
        "\n"
    )

    return structured_output

def main():
    st.title("Documents Summarization")

    text_input_from_user = st.text_area("Type or paste text directly here", height=300)
    uploaded_file = st.file_uploader("Or upload a file", type=["txt", "pdf", "docx", "csv"])

    if st.button("Generate Summary and Insights"):
        if text_input_from_user.strip():
            text_input = text_input_from_user
            st.success("Text input received successfully.")
        elif uploaded_file is not None:
            if uploaded_file.name.endswith('.txt'):
                text_input = uploaded_file.read().decode("utf-8")
                st.success(f"File uploaded successfully: {uploaded_file.name}")
            elif uploaded_file.name.endswith('.pdf'):
                pdf_reader = PdfReader(uploaded_file)
                text_input = "".join([page.extract_text() for page in pdf_reader.pages])
                st.success(f"PDF file uploaded successfully: {uploaded_file.name}")
            elif uploaded_file.name.endswith('.docx'):
                doc = Document(uploaded_file)
                text_input = "\n".join([para.text for para in doc.paragraphs])
                st.success(f"File uploaded successfully: {uploaded_file.name}")
            elif uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
                text_input = df.to_string(index=False)
                st.success(f"CSV file uploaded successfully: {uploaded_file.name}")
            else:
                st.error("Please upload a valid .txt, .pdf, .docx, or .csv file.")
                return
        else:
            st.error("Please either enter text or upload a file.")
            return

        if text_input.strip():
           structured_output = generate_structured_output(text_input)
           st.markdown(structured_output, unsafe_allow_html=True)


if __name__ == "__main__":
    main()
